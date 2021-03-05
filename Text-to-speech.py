#Creating a document with python
from docx import Document
from docx.shared import Inches
#Add text speech
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# Adding a profile picture
"""document.add_picture(
    "me.png",
    width = Inches(2.0)
)"""

# Name, Phone Number, Email
speak("What is your name?")
name = input("What is your name? ")
speak("Hello" + name + " Hope you are having a good day today")
speak("What is your phone number?")
phone_number = input("What is your phone number? ")
speak("Please enter your email address: ")
email_addy = input("Please enter your email address: ")

document.add_paragraph(
    name + " | " + phone_number + " | " + email_addy)

# About me
document.add_heading("About Me")
speak("Tell me about yourself")
document.add_paragraph(
    input("Tell me about yourself: ")
)
# work experience
document.add_heading("Work Experience")
p = document.add_paragraph()
speak("Tell me about your work history: ")
speak("Enter a company: ")
company = input("Enter Company: ")
speak("Enter you start and end date at " + company)
start_date = input("Start Date: ")
end_date = input("End Date: ")

p.add_run(company + " ").bold = True
p.add_run(start_date + "-" + end_date + "\n").italic = True
speak("Describe your experience at " + company + ":")
experience_detail = input("Describe your experience at " + company + ": ")
p.add_run(experience_detail)

# more work experience
while True:
    speak("Do you have more work experiences? Yes or No: ")

    has_more_experience = input("Do you have more work experience? Yes or No: ")
    if has_more_experience.lower() == "yes":
        p = document.add_paragraph()
        speak("Enter another company: ")
        company = input("Enter Company: ")
        speak("Enter you start and end date at " + company)
        start_date = input("Start Date: ")
        end_date = input("End Date: ")

        p.add_run(company + " ").bold = True
        p.add_run(start_date + "-" + end_date + "\n").italic = True
        speak("Describe your experience at " + company + ":")
        experience_detail = input("Describe your experience at " + company + ": ")
        p.add_run(experience_detail)
    else:
        break
#Skills
document.add_heading("Skills")
speak("Enter a skill: ")
skill = input("Enter a skill: ")
p = document.add_paragraph(skill)
p.style = "List Bullet"
while True:
    speak("Do you have another skill? Yes or No: ")
    has_more_skills = input("Do you have another skill? Yes or No: ")
    if has_more_skills.lower() == "yes":
        speak("Enter a skill: ")
        skill = input("Enter a skill: ")
        p = document.add_paragraph(skill)
        p.style = "List Bullet"
    else:
        break
#footer
"""section = document.sections[0]
footer = section.footer
p = footer.paragraph[0]
p.text = "CV generated using Amigoscode and Intuit Quickbooks project"
"""

document.save("cv.docx")


